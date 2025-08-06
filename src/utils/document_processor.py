import os
import logging
from typing import List, Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document
import pandas as pd
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangChainDocument

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )

    def process_file(self, file_path: str) -> List[LangChainDocument]:
        """Process uploaded file and return document chunks."""
        try:
            file_extension = Path(file_path).suffix.lower()
            if file_extension == '.pdf':
                text = self._extract_pdf_text(file_path)
            elif file_extension == '.docx':
                text = self._extract_docx_text(file_path)
            elif file_extension == '.txt':
                text = self._extract_txt_text(file_path)
            elif file_extension in ['.csv', '.xlsx']:
                text = self._extract_table_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")

            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Create LangChain documents
            documents = []
            for i, chunk in enumerate(chunks):
                doc = LangChainDocument(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "chunk_id": i,
                        "file_type": file_extension
                    }
                )
                documents.append(doc)
            
            logger.info(f"Processed {file_path}: {len(documents)} chunks created")
            return documents
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            raise

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()

    def _extract_table_text(self, file_path: str) -> str:
        """Extract text from CSV/XLSX files."""
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)
        
        # Convert dataframe to text representation
        text = f"Data from {Path(file_path).name}:\n"
        text += df.to_string(index=False)
        return text